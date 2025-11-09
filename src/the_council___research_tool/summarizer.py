import os
from pathlib import Path
from openai import AzureOpenAI

class FileSummarizer:
    def __init__(self):
        self.client = AzureOpenAI(
            api_version="2025-03-01-preview",
            azure_endpoint="https://nachalo-2324-resource.cognitiveservices.azure.com/",
            api_key="os.getenv("AZURE_OPENAI_API_KEY", "your-api-key-here")",
        )

    def summarize_file(self, file_path: str) -> str:
        """
        Summarize the content of a file using OpenAI with robust parsing and token management.

        Args:
            file_path: Path to the file to summarize

        Returns:
            Summary of the file content
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Read and parse file content based on file type
        try:
            content, metadata = self._parse_file_content(file_path)
        except Exception as e:
            raise ValueError(f"Cannot parse file content: {file_path} - {str(e)}")

        # Handle large files with chunking if needed
        if len(content) > 80000:  # ~80k characters to stay safe within token limits
            return self._summarize_large_file(content, metadata)

        # Create appropriate prompt based on file type
        prompt = self._create_summary_prompt(content, metadata)

        # Use Responses API with token limit handling
        try:
            response = self.client.responses.create(
                model="gpt-4o-mini",
                input=prompt,
                instructions="You are an expert document analyst. Provide clear, comprehensive summaries that capture key information, main arguments, and important details."
            )
            return response.output_text
        except Exception as e:
            if "token" in str(e).lower():
                # If still hitting token limits, truncate more aggressively
                truncated_content = content[:40000] + "\n\n[Content truncated for summarization]"
                prompt = self._create_summary_prompt(truncated_content, metadata)
                response = self.client.responses.create(
                    model="gpt-4o-mini",
                    input=prompt,
                    instructions="You are an expert document analyst. Provide clear, comprehensive summaries that capture key information, main arguments, and important details."
                )
                return response.output_text + "\n\n*Note: Summary based on truncated content due to document length.*"
            else:
                raise e

    def _parse_file_content(self, file_path: Path) -> tuple[str, dict]:
        """
        Parse file content based on file type with robust error handling.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (content, metadata)
        """
        file_extension = file_path.suffix.lower()
        metadata = {
            'filename': file_path.name,
            'file_type': file_extension,
            'file_size': file_path.stat().st_size
        }

        try:
            if file_extension in ['.pdf', '.docx', '.xlsx', '.xls']:
                content = self._parse_complex_file(file_path, file_extension)
            elif file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv']:
                content = self._parse_text_file(file_path)
            else:
                # Try to read as text anyway
                try:
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                    metadata['parsing_method'] = 'fallback_text'
                except:
                    raise ValueError(f"Unsupported file type: {file_extension}")

            metadata['content_length'] = len(content)
            return content, metadata

        except Exception as e:
            raise ValueError(f"File parsing failed: {str(e)}")

    def _parse_complex_file(self, file_path: Path, file_extension: str) -> str:
        """Parse complex file types (PDF, Word, Excel) using appropriate loaders with fallbacks."""
        content = ""

        try:
            if file_extension == '.pdf':
                content = self._parse_pdf_file(file_path)
            elif file_extension == '.docx':
                content = self._parse_docx_file(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                content = self._parse_excel_file(file_path)
            else:
                raise ValueError(f"No loader available for {file_extension}")

            return content

        except Exception as e:
            # Final fallback: try to read as plain text
            try:
                fallback_content = file_path.read_text(encoding='utf-8', errors='ignore')
                if fallback_content.strip():
                    return f"[Warning: Advanced parsing failed, using basic text extraction]\n\n{fallback_content}"
            except:
                pass

            raise ValueError(f"Failed to parse {file_extension} file: {str(e)}")

    def _parse_pdf_file(self, file_path: Path) -> str:
        """Parse PDF files with multiple fallback methods."""
        try:
            # Try PyPDFLoader first
            from langchain.document_loaders import PyPDFLoader
            loader = PyPDFLoader(str(file_path))
            documents = loader.load()
        except (ImportError, Exception):
            try:
                # Fallback to PyMuPDFLoader
                from langchain.document_loaders import PyMuPDFLoader
                loader = PyMuPDFLoader(str(file_path))
                documents = loader.load()
            except (ImportError, Exception):
                # Final fallback: basic text extraction
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(str(file_path))
                    content_parts = []
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        text = page.get_text()
                        if text.strip():
                            content_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                    doc.close()
                    return "\n\n".join(content_parts)
                except ImportError:
                    raise ValueError("PDF parsing requires 'pypdf' or 'PyMuPDF' library. Install with: pip install pypdf")

        # Extract content from loaded documents
        content_parts = []
        for i, doc in enumerate(documents):
            page_content = doc.page_content.strip()
            if page_content:
                if len(documents) > 1:
                    content_parts.append(f"--- Page {i+1} ---\n{page_content}")
                else:
                    content_parts.append(page_content)

        return "\n\n".join(content_parts)

    def _parse_docx_file(self, file_path: Path) -> str:
        """Parse Word documents with multiple fallback methods."""
        try:
            # Try Docx2txtLoader first
            from langchain.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(str(file_path))
            documents = loader.load()
        except (ImportError, Exception):
            try:
                # Fallback to python-docx
                import docx
                doc = docx.Document(str(file_path))
                content_parts = []

                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text)

                # Extract tables
                for table in doc.tables:
                    content_parts.append("\n--- Table ---\n")
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            content_parts.append(" | ".join(row_text))

                return "\n\n".join(content_parts)

            except ImportError:
                raise ValueError("Word document parsing requires 'docx2txt' or 'python-docx' library. Install with: pip install docx2txt or pip install python-docx")

        # Extract content from loaded documents
        content_parts = []
        for doc in documents:
            page_content = doc.page_content.strip()
            if page_content:
                content_parts.append(page_content)

        return "\n\n".join(content_parts)

    def _parse_excel_file(self, file_path: Path) -> str:
        """Parse Excel files with multiple fallback methods."""
        try:
            # Try UnstructuredExcelLoader first
            from langchain.document_loaders import UnstructuredExcelLoader
            loader = UnstructuredExcelLoader(str(file_path))
            documents = loader.load()
        except (ImportError, Exception):
            try:
                # Fallback to pandas
                import pandas as pd
                excel_data = pd.read_excel(str(file_path), sheet_name=None)

                content_parts = []
                for sheet_name, df in excel_data.items():
                    content_parts.append(f"--- Sheet: {sheet_name} ---")
                    content_parts.append(df.to_string(index=False))
                    content_parts.append("")  # Empty line between sheets

                return "\n\n".join(content_parts)

            except ImportError:
                raise ValueError("Excel parsing requires 'unstructured' or 'pandas' library. Install with: pip install unstructured or pip install pandas openpyxl")

        # Extract content from loaded documents
        content_parts = []
        for doc in documents:
            page_content = doc.page_content.strip()
            if page_content:
                content_parts.append(page_content)

        return "\n\n".join(content_parts)

    def _parse_text_file(self, file_path: Path) -> str:
        """Parse text-based files with encoding detection."""
        encodings_to_try = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings_to_try:
            try:
                content = file_path.read_text(encoding=encoding)
                # Basic validation - check if we got meaningful content
                if len(content.strip()) > 0:
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError("Could not decode file with any supported encoding")

    def _summarize_large_file(self, content: str, metadata: dict) -> str:
        """Handle summarization of large files by chunking and hierarchical summarization."""
        # Split content into chunks
        chunk_size = 30000  # ~30k characters per chunk
        chunks = []

        for i in range(0, len(content), chunk_size):
            chunk = content[i:i + chunk_size]
            if chunk.strip():  # Only add non-empty chunks
                chunks.append(chunk)

        if not chunks:
            return "No content found to summarize."

        # Summarize each chunk
        chunk_summaries = []
        for i, chunk in enumerate(chunks[:5]):  # Limit to first 5 chunks to avoid too many API calls
            try:
                chunk_prompt = f"Summarize this section of the document (Part {i+1}/{min(len(chunks), 5)}):\n\n{chunk[:25000]}"  # Extra truncation for chunks

                response = self.client.responses.create(
                    model="gpt-4o-mini",
                    input=chunk_prompt,
                    instructions="Provide a concise summary of this document section, focusing on key points and main ideas."
                )
                chunk_summaries.append(f"Section {i+1}: {response.output_text.strip()}")
            except Exception as e:
                chunk_summaries.append(f"Section {i+1}: [Summary failed: {str(e)}]")

        # Create final summary from chunk summaries
        final_prompt = f"""Create a comprehensive summary of this document from the following section summaries:

{' '.join(chunk_summaries)}

Document metadata: {metadata.get('filename', 'Unknown')} ({metadata.get('file_type', 'Unknown')}, {metadata.get('file_size', 0)} bytes)

Provide a cohesive summary that captures the main themes, key findings, and important information from across all sections."""

        try:
            response = self.client.responses.create(
                model="gpt-4o-mini",
                input=final_prompt,
                instructions="Create a unified, comprehensive summary that synthesizes information from all document sections."
            )
            return response.output_text + f"\n\n*Note: Summary based on {len(chunk_summaries)} document sections due to file size.*"
        except Exception as e:
            # Fallback: just return the chunk summaries
            return "Document Summary (by sections):\n\n" + "\n\n".join(chunk_summaries) + "\n\n*Note: Large document summarized by sections.*"

    def _create_summary_prompt(self, content: str, metadata: dict) -> str:
        """Create an appropriate summary prompt based on file type and content."""
        file_type = metadata.get('file_type', '').lower()
        filename = metadata.get('filename', 'document')

        # Customize prompt based on file type
        if file_type in ['.pdf']:
            prompt_type = "academic paper, report, or document"
        elif file_type in ['.docx']:
            prompt_type = "document or report"
        elif file_type in ['.xlsx', '.xls', '.csv']:
            prompt_type = "data table or spreadsheet"
        elif file_type in ['.py', '.js', '.html', '.css']:
            prompt_type = "code file"
        elif file_type in ['.json', '.xml']:
            prompt_type = "structured data file"
        elif file_type in ['.md']:
            prompt_type = "markdown document"
        else:
            prompt_type = "document"

        prompt = f"""Please provide a comprehensive summary of this {prompt_type}: "{filename}"

CONTENT:
{content}

SUMMARY REQUIREMENTS:
- Capture the main purpose and key information
- Include important findings, conclusions, or main points
- Note any key data, statistics, or technical details
- Maintain factual accuracy and preserve important context
- Keep the summary clear and well-structured
- For code files, explain the main functionality and purpose
- For data files, describe the structure and key insights

SUMMARY:"""

        return prompt

def main():
    import sys
    if len(sys.argv) != 2:
        print("Usage: python -m the_council___research_tool.summarizer <file_path>")
        sys.exit(1)

    file_path = sys.argv[1]
    summarizer = FileSummarizer()

    try:
        summary = summarizer.summarize_file(file_path)
        print("Summary:")
        print("=" * 50)
        print(summary)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
